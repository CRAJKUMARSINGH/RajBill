
import streamlit as st
import pandas as pd
import numpy as np
import os
import zipfile
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
import tempfile
import pdfkit
from PyPDF2 import PdfMerger
import base64
from jinja2 import Environment, FileSystemLoader
from pypdf import PdfReader, PdfWriter
import num2words
import platform
import traceback
from lxml import etree

# Temporary directory
TEMP_DIR = tempfile.mkdtemp()

# Configure wkhtmltopdf
if platform.system() == "Windows":
    wkhtmltopdf_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
else:
    config = pdfkit.configuration()

# Set up Jinja2 environment
env = Environment(loader=FileSystemLoader("templates"), cache_size=0)

# Helper functions
def number_to_words(number):
    try:
        return num2words(number, lang='en_IN')
    except:
        return str(number)

##########################################################################################
def merge_pdfs(pdf_files, output_path):
    merger = PdfMerger()
    for pdf in pdf_files:
        if os.path.exists(pdf):
            merger.append(pdf)
    merger.write(output_path)
    merger.close()

##########################################################################################
def process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type, amount_paid_last_bill, is_first_bill, user_inputs):
    first_page_data = {
        "header": [],
        "items": [],  # Initialize as empty list
        "totals": {}
    }
    last_page_data = {"payable_amount": 0, "amount_words": ""}
    deviation_data = {
        "items": [],  # Initialize as empty list
        "summary": {},
        "header": []
    }
    extra_items_data = {"items": []}
    note_sheet_data = {"notes": []}
################################################################################################################
    # Header (A1:I19)
    #header_data = ws_wo.iloc[:19].replace(np.nan, "").values.tolist()
    #first_page_data["header"] = header_data
    ###### REPLACEMENT 18 APRIL 2025
    from datetime import datetime, date

    # Header (A1:G19) — matching actual data range
    header_data = ws_wo.iloc[:19, :7].replace(np.nan, "").values.tolist()

    # Format dates in header_data
    for i in range(len(header_data)):
        for j in range(len(header_data[i])):
            val = header_data[i][j]
            if isinstance(val, (pd.Timestamp, datetime, date)):
                header_data[i][j] = val.strftime("%d-%m-%Y")

    # Assign header to both first_page_data and deviation_data
    first_page_data["header"] = header_data
    deviation_data["header"] = header_data  # Add to deviation_data
    
    # Ensure all dates are formatted as date-only strings (optional step, if needed before saving)
    for i in range(len(header_data)):
        for j in range(len(header_data[i])):
            val = header_data[i][j]
            if isinstance(val, (pd.Timestamp, datetime, date)):
                header_data[i][j] = val.strftime("%d-%m-%Y")

    # Assign to first page
    first_page_data["header"] = header_data
############################################################################################################
    # Work Order items
    last_row_wo = ws_wo.shape[0]
    for i in range(21, last_row_wo):
        qty_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else None
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                st.warning(f"Skipping invalid quantity at Bill Quantity row {i+1}: '{qty_raw}'")
                continue

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                st.warning(f"Skipping invalid rate at Work Order row {i+1}: '{rate_raw}'")
                continue

        item = {
            "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
            "quantity": qty,
            "rate": rate,
            "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
            "amount": round(qty * rate) if qty and rate else 0,
            "is_divider": False
        }
        first_page_data["items"].append(item)

    # Extra Items divider
    first_page_data["items"].append({
        "description": "Extra Items (With Premium)",
        "bold": True,
        "underline": True,
        "amount": 0,
        "quantity": 0,
        "rate": 0,
        "serial_no": "",
        "unit": "",
        "remark": "",
        "is_divider": True
    })

    # Extra Items
    last_row_extra = ws_extra.shape[0]
    for j in range(6, last_row_extra):
        qty_raw = ws_extra.iloc[j, 3] if pd.notnull(ws_extra.iloc[j, 3]) else None
        rate_raw = ws_extra.iloc[j, 5] if pd.notnull(ws_extra.iloc[j, 5]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                st.warning(f"Skipping invalid quantity at Extra Items row {j+1}: '{qty_raw}'")
                continue

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                st.warning(f"Skipping invalid rate at Extra Items row {j+1}: '{rate_raw}'")
                continue

        item = {
            "serial_no": str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else "",
            "description": str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else "",
            "unit": str(ws_extra.iloc[j, 4]) if pd.notnull(ws_extra.iloc[j, 4]) else "",
            "quantity": qty,
            "rate": rate,
            "remark": str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else "",
            "amount": round(qty * rate) if qty and rate else 0,
            "is_divider": False
        }
        first_page_data["items"].append(item)
        extra_items_data["items"].append(item.copy())  # Copy for standalone Extra Items

    # Totals
    data_items = [item for item in first_page_data["items"] if not item.get("is_divider", False)]
    total_amount = round(sum(item.get("amount", 0) for item in data_items))
    premium_amount = round(total_amount * (premium_percent / 100))
    payable_amount = round(total_amount + premium_amount)

    first_page_data["totals"] = {
        "grand_total": total_amount,
        "premium": {"percent": premium_percent / 100, "type": "above", "amount": premium_amount},
        "payable": payable_amount
    }

    try:
        extra_items_start = next(i for i, item in enumerate(first_page_data["items"]) if item.get("description") == "Extra Items (With Premium)")
        extra_items = [item for item in first_page_data["items"][extra_items_start + 1:] if not item.get("is_divider", False)]
        extra_items_sum = round(sum(item.get("amount", 0) for item in extra_items))
        extra_items_premium = round(extra_items_sum * (premium_percent / 100))
        first_page_data["totals"]["extra_items_sum"] = extra_items_sum + extra_items_premium
    except StopIteration:
        first_page_data["totals"]["extra_items_sum"] = 0

    # Last Page
    last_page_data = {"payable_amount": payable_amount, "amount_words": number_to_words(payable_amount)}

    # Deviation Statement
    work_order_total = 0
    executed_total = 0
    overall_excess = 0
    overall_saving = 0
    for i in range(21, last_row_wo):

        qty_wo_raw = ws_wo.iloc[i, 3] if pd.notnull(ws_wo.iloc[i, 3]) else None
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None
        qty_bill_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else None

        qty_wo = 0
        if isinstance(qty_wo_raw, (int, float)):
            qty_wo = float(qty_wo_raw)
        elif isinstance(qty_wo_raw, str):
            cleaned_qty_wo = qty_wo_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty_wo = float(cleaned_qty_wo)
            except ValueError:
                st.warning(f"Skipping invalid qty_wo at row {i+1}: '{qty_wo_raw}'")
                continue

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                st.warning(f"Skipping invalid rate at row {i+1}: '{rate_raw}'")
                continue

        qty_bill = 0
        if isinstance(qty_bill_raw, (int, float)):
            qty_bill = float(qty_bill_raw)
        elif isinstance(qty_bill_raw, str):
            cleaned_qty_bill = qty_bill_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty_bill = float(cleaned_qty_bill)
            except ValueError:
                st.warning(f"Skipping invalid qty_bill at row {i+1}: '{qty_bill_raw}'")
                continue

        amt_wo = round(qty_wo * rate)
        amt_bill = round(qty_bill * rate)
        excess_qty = qty_bill - qty_wo if qty_bill > qty_wo else 0
        excess_amt = round(excess_qty * rate) if excess_qty > 0 else 0
        saving_qty = qty_wo - qty_bill if qty_bill < qty_wo else 0
        saving_amt = round(saving_qty * rate) if saving_qty > 0 else 0

        item = {
            "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
            "qty_wo": qty_wo,
            "rate": rate,
            "amt_wo": amt_wo,
            "qty_bill": qty_bill,
            "amt_bill": amt_bill,
            "excess_qty": excess_qty,
            "excess_amt": excess_amt,
            "saving_qty": saving_qty,
            "saving_amt": saving_amt
        }
        deviation_data["items"].append(item)
        work_order_total += amt_wo
        executed_total += amt_bill
        overall_excess += excess_amt
        overall_saving += saving_amt

    # Deviation Summary
    tender_premium_f = round(work_order_total * (premium_percent / 100))
    tender_premium_h = round(executed_total * (premium_percent / 100))
    tender_premium_j = round(overall_excess * (premium_percent / 100))
    tender_premium_l = round(overall_saving * (premium_percent / 100))
    grand_total_f = work_order_total + tender_premium_f
    grand_total_h = executed_total + tender_premium_h
    grand_total_j = overall_excess + tender_premium_j
    grand_total_l = overall_saving + tender_premium_l
    net_difference = grand_total_h - grand_total_f

    deviation_data["summary"] = {
        "work_order_total": round(work_order_total),
        "executed_total": round(executed_total),
        "overall_excess": round(overall_excess),
        "overall_saving": round(overall_saving),
        "premium": {"percent": premium_percent / 100, "type": "above"},
        "tender_premium_f": tender_premium_f,
        "tender_premium_h": tender_premium_h,
        "tender_premium_j": tender_premium_j,
        "tender_premium_l": tender_premium_l,
        "grand_total_f": grand_total_f,
        "grand_total_h": grand_total_h,
        "grand_total_j": grand_total_j,
        "grand_total_l": grand_total_l,
        "net_difference": round(net_difference)
    }


    # Calculate work order amount
    try:
        work_order_amount = sum(
            float(ws_wo.iloc[i, 3]) * float(ws_wo.iloc[i, 4])
            for i in range(21, ws_wo.shape[0])
            if pd.notnull(ws_wo.iloc[i, 3]) and pd.notnull(ws_wo.iloc[i, 4])
        )
    except Exception as e:
        st.error(f"Error calculating work_order_amount: {e}")
        work_order_amount = 854678  # Fallback value

    # Calculate totals for certificate_iii
    extra_item_amount = first_page_data["totals"].get("extra_items_sum", 0)
    payable_amount = first_page_data["totals"].get("payable", 0)
    total_123 = payable_amount + extra_item_amount

    # Prepare data for certificate_iii
    certificate_iii_data = {
        "payable_amount": payable_amount,
        "total_123": total_123,
        "balance_4_minus_5": payable_amount,
        "amount_paid_last_bill": 0,  # Default value for first bill
        "payment_now": payable_amount,
        "by_cheque": payable_amount,
        "cheque_amount_words": number_to_words(payable_amount),
        "certificate_items": [
            {"name": "Total value of work", "percentage": "100%", "value": total_123},
            {"name": "Less: Amount Paid Last Bill", "percentage": "-", "value": 0},  # Default value for first bill
            {"name": "Net Payable", "percentage": "-", "value": payable_amount}
        ],
        "total_recovery": 0,  # Add logic for recovery items if needed
        "totals": {
            "grand_total": total_123,
            "payable_amount": payable_amount,
            "extra_items_sum": extra_item_amount,
            "total_123": total_123,
            "grand_total": total_123
        }
    }

    # Prepare Certificate III data
    certificate_iii_data = {
        'payable_amount': payable_amount,
        'amount_words': number_to_words(payable_amount),
        'totals': {
            'grand_total': first_page_data['totals']['grand_total'],
            'payable_amount': payable_amount,
            'extra_items_sum': first_page_data['totals']['extra_items_sum'],
            'total_123': first_page_data['totals']['grand_total']
        },
        'total_123': first_page_data['totals']['grand_total'],
        'balance_4_minus_5': first_page_data['totals']['grand_total'],
        'amount_paid_last_bill': 0,  # Default value for first bill
        'payment_now': payable_amount,
        'by_cheque': payable_amount,
        'cheque_amount_words': number_to_words(payable_amount),
        'certificate_items': [
            {
                'name': 'Total value of work',
                'percentage': '100%',
                'value': first_page_data['totals']['grand_total']
            },
            {
                'name': 'Less: Amount Paid Last Bill',
                'percentage': '-',
                'value': 0  # Default value for first bill
            },
            {
                'name': 'Net Payable',
                'percentage': '-',
                'value': payable_amount
            }
        ],
        'total_recovery': 0,
        'measurement_officer': 'Measurement Officer Name',
        'measurement_date': '30/04/2025',
        'measurement_book_page': '123',
        'measurement_book_no': 'MB-001',
        'officer_name': 'Officer Name',
        'officer_designation': 'Designation',
        'authorising_officer_name': 'Authorising Officer Name',
        'authorising_officer_designation': 'Designation'
    }

    # Generate note sheet
    note_sheet_data = {
        'agreement_no': ws_wo.iloc[0, 1] if pd.notnull(ws_wo.iloc[0, 1]) else '48/2024-25',
        'name_of_work': ws_wo.iloc[1, 1] if pd.notnull(ws_wo.iloc[1, 1]) else 'Electric Repair and MTC work at Govt. Ambedkar hostel Ambamata, Govardhanvilas, Udaipur',
        'name_of_firm': ws_wo.iloc[2, 1] if pd.notnull(ws_wo.iloc[2, 1]) else 'M/s Seema Electrical Udaipur',
        'date_commencement': ws_wo.iloc[3, 1] if pd.notnull(ws_wo.iloc[3, 1]) else '18/01/2025',
        'date_completion': ws_wo.iloc[4, 1] if pd.notnull(ws_wo.iloc[4, 1]) else '17/04/2025',
        'actual_completion': ws_wo.iloc[5, 1] if pd.notnull(ws_wo.iloc[5, 1]) else '01/03/2025',
        'work_order_amount': str(work_order_amount),
        'extra_item_amount': extra_item_amount,
        'notes': generate_bill_notes(payable_amount, work_order_amount, extra_item_amount),
        'totals': {
            'grand_total': first_page_data['totals']['grand_total'],
            'premium': first_page_data['totals']['premium'],
            'payable': first_page_data['totals']['payable'],
            'extra_items_sum': first_page_data['totals']['extra_items_sum']
        }
    }

    return first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data, certificate_iii_data

##########################################################################################
def generate_bill_notes(payable_amount, work_order_amount, extra_item_amount):
    percentage_work_done = float(payable_amount / work_order_amount * 100) if work_order_amount > 0 else 0
    serial_number = 1
    note = []
    note.append(f"{serial_number}. The work has been completed {percentage_work_done:.2f}% of the Work Order Amount.")
    serial_number += 1
    if percentage_work_done < 90:
        note.append(f"{serial_number}. The execution of work at final stage is less than 90%...")
        serial_number += 1
    elif percentage_work_done > 100 and percentage_work_done <= 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed...")
        serial_number += 1
    elif percentage_work_done > 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed...")
        serial_number += 1
    note.append(f"{serial_number}. Quality Control (QC) test reports attached.")
    serial_number += 1
    if extra_item_amount > 0:
        extra_item_percentage = float(extra_item_amount / work_order_amount * 100) if work_order_amount > 0 else 0
        if extra_item_percentage > 5:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}...")
        else:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}...")
        serial_number += 1
    note.append(f"{serial_number}. Please peruse above details for necessary decision-making.")
    note.append("")
    note.append("                                Premlata Jain")
    note.append("                               AAO- As Auditor")
    return {
        "notes": note,
        "work_order_amount": work_order_amount,
        "totals": {
            "grand_total": payable_amount,
            "premium": {"percent": 0.0, "type": "above", "amount": 0},
            "payable": payable_amount,
            "extra_items_sum": extra_item_amount
        }
    }

##########################################################################################
def generate_pdf(sheet_name, data, orientation, output_path):
    try:
        # Define required fields based on sheet name
        required_fields = {
            "First Page": [
                "header", "items", "totals"
            ],
            "Last Page": [
                "payable_amount", "amount_words"
            ],
            "Deviation Statement": [
                "items", "summary", "header"
            ],
            "Extra Items": [
                "items"
            ],
            "Note Sheet": [
                "agreement_no", "name_of_work", "name_of_firm",
                "date_commencement", "date_completion", "actual_completion",
                "work_order_amount", "extra_item_amount", "notes", "totals"
            ],
            "Certificate III": [
                "payable_amount", "total_123", "balance_4_minus_5",
                "amount_paid_last_bill", "payment_now", "by_cheque",
                "cheque_amount_words", "certificate_items",
                "total_recovery", "totals"
            ]
        }

        # Get required fields for this sheet
        required = required_fields.get(sheet_name, [])
        
        # Validate required fields
        for field in required:
            if field not in data:
                raise ValueError(f"Missing required field for {sheet_name}: {field}")
                
        # Validate totals dictionary if it exists
        if "totals" in required and "totals" in data:
            required_totals = {
                "First Page": [
                    "grand_total", "premium", "payable"
                ],
                "Certificate III": [
                    "grand_total", "payable_amount", "extra_items_sum",
                    "total_123"
                ]
            }
            required = required_totals.get(sheet_name, [])
            for field in required:
                if field not in data["totals"]:
                    raise ValueError(f"Missing required totals field for {sheet_name}: {field}")

        template = env.get_template(f"{sheet_name.lower().replace(' ', '_')}.html")
        html_content = template.render(data=data)
        
        # Save HTML for debugging
        debug_html_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}_debug.html")
        with open(debug_html_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        # Base options for all PDFs
        options = {
            "page-size": "A4",
            "orientation": orientation,
            "margin-top": "0.25in",
            "margin-bottom": "0.25in",
            "margin-left": "0.25in",
            "margin-right": "0.5in",
            "encoding": "UTF-8",
            "quiet": "",
            "no-outline": None,
            "enable-local-file-access": None,
            "disable-smart-shrinking": None,
            "dpi": 300,
            "javascript-delay": "1000",
            "no-stop-slow-scripts": None,
            "load-error-handling": "ignore",
            "debug-javascript": None,
            "disable-external-links": None,
            "disable-internal-links": None,
            "disable-forms": None,
            "disable-smart-shrinking": None,
            "disable-javascript": None,
            "disable-local-file-access": None
        }
        
        if sheet_name == "Note Sheet":
            options["margin-bottom"] = "0.6in"
        elif sheet_name == "Deviation Statement":
            options["margin-bottom"] = "0.25in"

        try:
            # Create parent directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Generate PDF
            pdfkit.from_string(
                html_content,
                output_path,
                configuration=config,
                options=options
            )
            
            # Verify PDF was created
            if not os.path.exists(output_path):
                raise Exception(f"PDF file was not created at {output_path}")
                
            return True
            
        except Exception as e:
            st.error(f"Error generating PDF for {sheet_name}: {str(e)}")
            st.write(traceback.format_exc())
            raise

    except Exception as e:
        st.error(f"Error generating PDF for {sheet_name}: {str(e)}")
        st.write(traceback.format_exc())
        raise

##########################################################################################
def create_word_doc(sheet_name, data, doc_path, orientation="portrait"):
    try:
        doc = Document()
        # Set page orientation
        section = doc.sections[0]
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        
        # Set margins (14mm left/right, 10mm bottom for First Page)
        section.left_margin = Inches(0.5512)  # 14mm
        section.right_margin = Inches(0.5512)
        section.top_margin = Inches(0.5512)
        section.bottom_margin = Inches(0.3937)  # 10mm

        if sheet_name == "First Page":
            # Add header
            doc.add_heading("CONTRACTOR BILL", level=2)
            
            # Safely add header rows from data
            if "header" in data and isinstance(data["header"], list):
                for row in data["header"]:
                    if isinstance(row, (list, tuple)) and row:  # Check if row is a non-empty sequence
                        p = doc.add_paragraph()
                        p.add_run(" ".join(str(cell) for cell in row if str(cell).strip())).font.size = Pt(8)

            # Only proceed with table creation if items exist
            if "items" in data and isinstance(data["items"], list) and data["items"]:
                # Create table with appropriate rows
                table_rows = 2 + len(data["items"]) + 3
                table = doc.add_table(rows=table_rows, cols=9)
                # Add custom table style with borders
                table.style = "Table Grid"
                
                # Add borders to all cells
                for row in table.rows:
                    for cell in row.cells:
                        # Set font size for cell text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        # Add borders to cell
                        set_cell_border(cell, 
                            top={"val": "single", "sz": "6", "color": "000000"},
                            bottom={"val": "single", "sz": "6", "color": "000000"},
                            left={"val": "single", "sz": "6", "color": "000000"},
                            right={"val": "single", "sz": "6", "color": "000000"}
                        )
                
                # Set column widths
                column_widths = [
                    0.4,  # Unit
                    0.55,  # Quantity Since Last
                    0.55,  # Quantity Upto Date
                    0.38,  # Serial No
                    2.5,  # Description
                    0.52,  # Rate
                    0.77,  # Amount
                    0.6,  # Amount Previous
                    0.47   # Remark
                ]
                for col_idx, width in enumerate(column_widths):
                    for cell in table.columns[col_idx].cells:
                        cell.width = Inches(width)

                # Add table header
                table_header = [
                    "Unit",
                    "Quantity executed (or supplied) since last certificate",
                    "Quantity executed (or supplied) upto date as per MB",
                    "S. No.",
                    "Item of Work supplies (Grouped under \"sub-head\" and \"sub work\" of estimate)",
                    "Rate",
                    "Upto date Amount",
                    "Amount Since previous bill (Total for each sub-head)",
                    "Remarks"
                ]
                header_row = table.rows[0]
                for i, header in enumerate(table_header):
                    header_row.cells[i].text = header
                    header_row.cells[i].paragraphs[0].runs[0].font.size = Pt(7)

                # Add header numbers (1-9)
                number_row = table.rows[1]
                for i in range(9):
                    number_row.cells[i].text = str(i + 1)
                    number_row.cells[i].paragraphs[0].runs[0].font.size = Pt(7)

                # Add data rows
                for i, item in enumerate(data["items"]):
                    if i + 2 >= len(table.rows):
                        break  # Prevent index errors
                    row = table.rows[i + 2]
                    row.cells[0].text = str(item.get("unit", ""))
                    row.cells[1].text = str(item.get("quantity_since_last", ""))
                    row.cells[2].text = str(item.get("quantity_upto_date", item.get("quantity", "")))
                    row.cells[3].text = str(item.get("serial_no", ""))
                    row.cells[4].text = str(item.get("description", ""))
                    row.cells[5].text = str(item.get("rate", ""))
                    row.cells[6].text = str(item.get("amount", ""))
                    row.cells[7].text = str(item.get("amount_previous", ""))
                    row.cells[8].text = str(item.get("remark", ""))
                    
                    # Apply styling to Description
                    for cell in row.cells:
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if item.get("bold"):
                        row.cells[4].paragraphs[0].runs[0].font.bold = True
                    if item.get("underline"):
                        row.cells[4].paragraphs[0].runs[0].font.underline = True

                # Add totals rows if they exist in the table
                if "totals" in data and len(table.rows) >= 3:
                    if len(table.rows) > len(data["items"]) + 2:
                        totals_row = table.rows[-3]
                        totals_row.cells[4].text = "Grand Total"
                        totals_row.cells[4].paragraphs[0].runs[0].font.bold = True
                        totals_row.cells[6].text = str(data["totals"].get("grand_total", ""))
                        totals_row.cells[6].paragraphs[0].runs[0].font.bold = True

                    if len(table.rows) > len(data["items"]) + 1:
                        premium_row = table.rows[-2]
                        premium_percent = data["totals"].get("premium", {}).get("percent", 0) * 100
                        premium_row.cells[4].text = f"Premium ({premium_percent:.2f}% {data['totals']['premium']['type']})"
                        premium_row.cells[4].paragraphs[0].runs[0].font.bold = True
                        premium_row.cells[6].text = str(data["totals"].get("premium", {}).get("amount", ""))
                        premium_row.cells[6].paragraphs[0].runs[0].font.bold = True

                    if len(table.rows) > len(data["items"]):
                        payable_row = table.rows[-1]
                        payable_row.cells[4].text = "Total Payable"
                        payable_row.cells[4].paragraphs[0].runs[0].font.bold = True
                        payable_row.cells[6].text = str(data["totals"].get("payable", ""))
                        payable_row.cells[6].paragraphs[0].runs[0].font.bold = True

                # Add deviation statement if it exists in the data
                if "deviation" in data and data["deviation"]:
                    # Add spacing paragraph
                    doc.add_paragraph()
                    
                    # Add deviation header
                    deviation_header = doc.add_heading("DEVIATION STATEMENT", level=3)
                    deviation_header.paragraph_format.space_after = Pt(12)
                    
                    # Create deviation table
                    deviation_table = doc.add_table(rows=1, cols=4)
                    deviation_table.style = "Table Grid"
                    
                    # Set column widths
                    deviation_column_widths = [1.5, 1.5, 1.5, 1.5]
                    for col_idx, width in enumerate(deviation_column_widths):
                        for cell in deviation_table.columns[col_idx].cells:
                            cell.width = Inches(width)
                    
                    # Add headers
                    headers = ["Item", "Work Order Qty", "Executed Qty", "Deviation"]
                    for i, header in enumerate(headers):
                        deviation_table.cell(0, i).text = header
                        deviation_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
                    
                    # Add deviation data
                    deviation_items = data["deviation"]["items"]
                    for i, item in enumerate(deviation_items):
                        if i >= len(deviation_table.rows) - 1:
                            deviation_table.add_row()
                        row = deviation_table.rows[i + 1]
                        row.cells[0].text = str(item.get("description", ""))
                        row.cells[1].text = str(item.get("work_order_qty", ""))
                        row.cells[2].text = str(item.get("executed_qty", ""))
                        row.cells[3].text = str(item.get("deviation", ""))
                        
                        # Apply styling
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)
                    
                    # Add summary
                    if "summary" in data["deviation"]:
                        summary = data["deviation"]["summary"]
                        summary_text = f"\nOverall Excess: {summary.get('excess', 0)}\nOverall Saving: {summary.get('saving', 0)}"
                        summary_paragraph = doc.add_paragraph(summary_text)
                        summary_paragraph.paragraph_format.space_before = Pt(12)
                        for run in summary_paragraph.runs:
                            run.font.size = Pt(8)
                            run.font.bold = True

        elif sheet_name == "Last Page":
            doc.add_heading("PAYABLE AMOUNT", level=2)
            p = doc.add_paragraph()
            p.add_run(f"Payable Amount: {data.get('payable_amount', '')}").font.size = Pt(8)
            p = doc.add_paragraph()
            p.add_run(f"Total in Words: {data.get('amount_words', '')}").font.size = Pt(8)
            
        elif sheet_name == "Extra Items":
            if "items" in data and isinstance(data["items"], list) and data["items"]:
                doc.add_heading("EXTRA ITEMS", level=2)
                table = doc.add_table(rows=len(data["items"]) + 1, cols=7)
                table.style = "Table Grid"
                
                # Add borders to all cells
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        set_cell_border(cell, 
                            top={"val": "single", "sz": "6", "color": "000000"},
                            bottom={"val": "single", "sz": "6", "color": "000000"},
                            left={"val": "single", "sz": "6", "color": "000000"},
                            right={"val": "single", "sz": "6", "color": "000000"}
                        )
                
                headers = ["Serial No.", "Remark", "Description", "Quantity", "Unit", "Rate", "Amount"]
                for j, header in enumerate(headers):
                    table.rows[0].cells[j].text = header
                    table.rows[0].cells[j].paragraphs[0].runs[0].font.size = Pt(7)
                
                for i, item in enumerate(data["items"]):
                    if i + 1 < len(table.rows):
                        row = table.rows[i + 1]
                        row.cells[0].text = str(item.get("serial_no", ""))
                        row.cells[1].text = str(item.get("remark", ""))
                        row.cells[2].text = str(item.get("description", ""))
                        row.cells[3].text = str(item.get("quantity", ""))
                        row.cells[4].text = str(item.get("unit", ""))
                        row.cells[5].text = str(item.get("rate", ""))
                        row.cells[6].text = str(item.get("amount", ""))
                        
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
        elif sheet_name == "Deviation Statement":
            # Set landscape orientation for Deviation Statement
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            
            # Set margins for landscape
            section.left_margin = Inches(0.5512)  # 14mm
            section.right_margin = Inches(0.5512)
            section.top_margin = Inches(0.5512)
            section.bottom_margin = Inches(0.3937)  # 10mm

            if "items" in data and isinstance(data["items"], list) and data["items"]:
                doc.add_heading("DEVIATION STATEMENT", level=2)
                table = doc.add_table(rows=2 + len(data["items"]), cols=8)
                table.style = "Table Grid"
                
                # Add borders to all cells
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        set_cell_border(cell, 
                            top={"val": "single", "sz": "6", "color": "000000"},
                            bottom={"val": "single", "sz": "6", "color": "000000"},
                            left={"val": "single", "sz": "6", "color": "000000"},
                            right={"val": "single", "sz": "6", "color": "000000"}
                        )
                
                # Add table headers
                headers = [
                    "S. No.",
                    "Description of Work", 
                    "Quantity as per Bill", 
                    "Quantity as per MB", 
                    "Quantity as per MB", 
                    "Quantity as per MB", 
                    "Quantity as per MB", 
                    "Remarks"
                ]
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                    table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(7)
                
                # Add header numbers
                for i in range(8):
                    table.rows[1].cells[i].text = str(i + 1)
                    table.rows[1].cells[i].paragraphs[0].runs[0].font.size = Pt(7)
                
                # Add data rows
                for i, item in enumerate(data["items"]):
                    if i + 2 >= len(table.rows):
                        break
                    row = table.rows[i + 2]
                    row.cells[0].text = str(item.get("serial_no", ""))
                    row.cells[1].text = str(item.get("description", ""))
                    row.cells[2].text = str(item.get("quantity_bill", ""))
                    row.cells[3].text = str(item.get("quantity_mb", ""))
                    row.cells[4].text = str(item.get("quantity_mb", ""))
                    row.cells[5].text = str(item.get("quantity_mb", ""))
                    row.cells[6].text = str(item.get("quantity_mb", ""))
                    row.cells[7].text = str(item.get("remark", ""))
                    
                    for cell in row.cells:
                        cell.paragraphs[0].runs[0].font.size = Pt(8)

        # Save the document
        doc.save(doc_path)
        return True

    except Exception as e:
        st.error(f"Error generating Word document for {sheet_name}: {str(e)}")
        st.write(traceback.format_exc())
        raise

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": "000000"},
        bottom={"val": "single", "sz": "6", "color": "000000"},
        left={"val": "single", "sz": "6", "color": "000000"},
        right={"val": "single", "sz": "6", "color": "000000"}
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Get or create the tcBorders element
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # Add or update each border
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            # Create or find the specific border element
            border = tcBorders.find(qn(f'w:{edge}'))
            if border is None:
                border = OxmlElement(f'w:{edge}')
                tcBorders.append(border)
            
            # Set the border attributes
            for key, value in edge_data.items():
                border.set(qn(f'w:{key}'), str(value))

def qn(tag):
    """Return a qualified name (with namespace) for the given tag."""
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }
    prefix, tagname = tag.split(':')
    return '{%s}%s' % (nsmap[prefix], tagname)

##########################################################################################
def main():
    # Add custom CSS for styling
    st.markdown("""
    <style>
        .stButton > button {
            background-color: #4CAF50;
            color: white;
            padding: 10px 20px;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 16px;
        }
        .stButton > button:hover {
            background-color: #45a049;
        }
        .stFileUploader > div > div {
            background-color: #f5f5f5;
            padding: 10px;
            border-radius: 4px;
            border: 1px solid #ddd;
        }
        .stFileUploader > div > div:hover {
            background-color: #e0e0e0;
        }
        .stFormSubmitButton > button {
            background-color: #2196F3 !important;
            color: white !important;
        }
        .stFormSubmitButton > button:hover {
            background-color: #1976D2 !important;
        }
    </style>
    """, unsafe_allow_html=True)

    # Application description
    st.markdown("""
    <div style='text-align: center; padding: 20px; background-color: #f9f9f9; border-radius: 8px;'>
        <h3>Generate Contractor Bills with Ease</h3>
        <p>Upload your Excel files and generate professional contractor bills in seconds.</p>
    </div>
    """, unsafe_allow_html=True)

    # Main content
    st.markdown("""
    ### Instructions:
    1. Fill in the required details in the sidebar
    2. Upload an Excel file containing three sheets:
       - Work Order (ws_wo)
       - Bill Quantity (ws_bq)
       - Extra Items (ws_extra)
    """, unsafe_allow_html=True)

    # File upload
    uploaded_file = st.file_uploader("Upload Excel File", type=["xlsx", "xls"])

    # Get user inputs from sidebar
    premium_percent = st.sidebar.number_input(
        "Premium Percentage",
        min_value=0.0,
        max_value=100.0,
        value=0.0,
        step=0.1
    )
    premium_type = st.sidebar.selectbox(
        "Premium Type",
        ["Above", "Below"]
    )
    amount_paid_last_bill = st.sidebar.number_input(
        "Amount Paid Last Bill",
        min_value=0,
        value=0
    )
    is_first_bill = st.sidebar.checkbox("Is First Bill")

    if st.button("Generate Bill"):
        if uploaded_file is not None:
            try:
                # Read the uploaded file
                with pd.ExcelFile(uploaded_file) as xls:
                    ws_wo = pd.read_excel(xls, "Work Order", header=None)
                    ws_bq = pd.read_excel(xls, "Bill Quantity", header=None)
                    ws_extra = pd.read_excel(xls, "Extra Items", header=None)

                # Process the bill
                first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data, certificate_iii_data = process_bill(
                    ws_wo,
                    ws_bq,
                    ws_extra,
                    premium_percent,
                    premium_type,
                    amount_paid_last_bill,
                    is_first_bill,
                    user_inputs={
                        "premium_percent": premium_percent,
                        "premium_type": premium_type,
                        "amount_paid_last_bill": amount_paid_last_bill,
                        "is_first_bill": is_first_bill
                    }
                )

                # Generate PDFs
                pdf_files = []
                for sheet_name, data, orientation, template_name in [
                    ("First Page", first_page_data, "portrait", "first_page"),
                    ("Last Page", last_page_data, "portrait", "last_page"),
                    ("Extra Items", extra_items_data, "portrait", "extra_items"),
                    ("Deviation Statement", deviation_data, "landscape", "deviation_statement"),
                    ("Note Sheet", note_sheet_data, "portrait", "note_sheet"),
                    ("Certificate III", certificate_iii_data, "portrait", "certificate_iii")
                ]:
                    pdf_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}.pdf")
                    if generate_pdf(template_name, data, orientation, pdf_path):
                        pdf_files.append(pdf_path)

                # Merge PDFs
                current_date = datetime.now().strftime("%Y%m%d")
                pdf_output = os.path.join(TEMP_DIR, f"BILL_AND_DEVIATION_{current_date}.pdf")
                merge_pdfs(pdf_files, pdf_output)

                # Generate Word documents
                word_files = []
                for sheet_name, data in [
                    ("First Page", first_page_data),
                    ("Last Page", last_page_data),
                    ("Extra Items", extra_items_data),
                    ("Deviation Statement", deviation_data),
                    ("Note Sheet", note_sheet_data),
                    ("Certificate III", certificate_iii_data)
                ]:
                    doc_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}.docx")
                    if create_word_doc(sheet_name, data, doc_path):
                        word_files.append(doc_path)

                # Create ZIP file
                zip_path = os.path.join(TEMP_DIR, "output.zip")
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                    zipf.write(pdf_output, os.path.basename(pdf_output))
                    for word_file in word_files:
                        zipf.write(word_file, os.path.basename(word_file))

                # Provide download link
                with open(zip_path, "rb") as f:
                    bytes_data = f.read()
                st.download_button(
                    label="Download Output Files",
                    data=bytes_data,
                    file_name="bill_output.zip",
                    mime="application/zip"
                )
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
                st.stop()

if __name__ == "__main__":
    main()