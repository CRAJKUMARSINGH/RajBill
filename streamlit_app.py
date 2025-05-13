import streamlit as st
import pandas as pd
import pdfkit
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
import os
import shutil
from datetime import datetime, date
import zipfile
import tempfile
from PyPDF2 import PdfMerger
import concurrent.futures
from functools import lru_cache, partial
from typing import Dict, List, Tuple, Union, Any, Callable
import logging
import traceback
import platform
from jinja2 import Environment, FileSystemLoader
from pypdf import PdfReader, PdfWriter
import platform
import traceback
from lxml import etree

# Temporary directory
TEMP_DIR = tempfile.mkdtemp()

# Configure wkhtmltopdf
if platform.system() == "Windows":
    wkhtmltopdf_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
else:
    config = pdfkit.configuration()

# Set up Jinja2 environment
env = Environment(loader=FileSystemLoader("templates"), cache_size=0)

# Helper functions
def number_to_words(number):
    try:
        return num2words(number, lang='en_IN')
    except:
        return str(number)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('bill_generator.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Create a temporary directory
TEMP_DIR = tempfile.mkdtemp()

# Set the path to wkhtmltopdf
path_wkhtmltopdf = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe' # raw string is important here.
config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)

class BillGenerationError(Exception):
    """Custom exception for bill generation errors"""
    pass

def handle_error(error: Exception, context: str = "") -> None:
    """
    Handle errors consistently across the application.
    
    Args:
        error: The exception that occurred
        context: Additional context about where the error occurred
    """
    error_msg = f"Error in {context}: {str(error)}"
    logger.error(error_msg)
    logger.error(traceback.format_exc())
    st.error(error_msg)
    raise BillGenerationError(error_msg)

def cleanup_temp_files() -> None:
    """Clean up temporary files and directories"""
    try:
        if os.path.exists(TEMP_DIR):
            shutil.rmtree(TEMP_DIR)
            logger.info("Temporary files cleaned up successfully")
    except Exception as e:
        handle_error(e, "cleanup_temp_files")

def validate_excel_sheets(xls: pd.ExcelFile) -> None:
    """
    Validate the Excel file structure.
    
    Args:
        xls: The Excel file to validate
        
    Raises:
        ValueError: If the Excel file is invalid
    """
    try:
        required_sheets = ["Work Order", "Bill Quantity", "Extra Items"]
        missing_sheets = [sheet for sheet in required_sheets if sheet not in xls.sheet_names]
        if missing_sheets:
            raise ValueError(f"Missing required sheets: {', '.join(missing_sheets)}")
        
        # Validate sheet structure
        for sheet_name in required_sheets:
            df = pd.read_excel(xls, sheet_name, header=None)
            if df.empty:
                raise ValueError(f"Sheet '{sheet_name}' is empty")
            if sheet_name == "Work Order" and df.shape[1] < 7:
                raise ValueError("Work Order sheet must have at least 7 columns")
            if sheet_name == "Bill Quantity" and df.shape[1] < 4:
                raise ValueError("Bill Quantity sheet must have at least 4 columns")
            if sheet_name == "Extra Items" and df.shape[1] < 6:
                raise ValueError("Extra Items sheet must have at least 6 columns")
                
        logger.info("Excel file validation successful")
    except Exception as e:
        handle_error(e, "validate_excel_sheets")

@lru_cache(maxsize=128)
def number_to_words(number: Union[int, float]) -> str:
    """
    Convert a number to words in Indian number system.
    Cached for better performance.
    
    Args:
        number: The number to convert to words
        
    Returns:
        str: The number in words (Indian system)
        
    Raises:
        ValueError: If the number is negative
    """
    if number < 0:
        raise ValueError("Number must be non-negative")
    return num2words(int(number), lang='en_IN').title()

def process_bill_items_parallel(items: List[Dict[str, Any]], process_func: Callable) -> List[Dict[str, Any]]:
    """
    Process bill items in parallel for better performance.
    
    Args:
        items: List of items to process
        process_func: Function to process each item
        
    Returns:
        List of processed items
    """
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = [executor.submit(process_func, item) for item in items]
        return [future.result() for future in concurrent.futures.as_completed(futures)]

def process_bill(
    ws_wo: pd.DataFrame,
    ws_bq: pd.DataFrame,
    ws_extra: pd.DataFrame,
    premium_percent: float,
    premium_type: str,
    amount_paid_last_bill: float,
    is_first_bill: bool,
    user_inputs: Dict[str, Any]
) -> Tuple[Dict[str, Any], Dict[str, Any], Dict[str, Any], List[Dict[str, Any]], Dict[str, Any], Dict[str, Any]]:
    """
    Process bill data and generate all required documents.
    Optimized for performance with parallel processing.
    
    Args:
        ws_wo: Work order DataFrame
        ws_bq: Bill quantity DataFrame
        ws_extra: Extra items DataFrame
        premium_percent: Premium percentage (with 2 decimal places)
        premium_type: Type of premium ('above' or 'below')
        amount_paid_last_bill: Amount paid in last bill (integer)
        is_first_bill: Whether this is the first bill
        user_inputs: Dictionary containing user inputs
        
    Returns:
        Tuple containing:
        - First page data
        - Last page data
        - Deviation statement data
        - Extra items data
        - Note sheet data
        - Certificate III data
        
    Raises:
        ValueError: If input validation fails
        Exception: If processing fails
    """
    try:
        # Input validation
        if not isinstance(ws_wo, pd.DataFrame) or not isinstance(ws_bq, pd.DataFrame):
            raise ValueError("Invalid input data: work order and bill quantity must be pandas DataFrames")
        
        if not isinstance(premium_percent, (int, float)) or premium_percent < 0:
            raise ValueError("Premium percent must be a non-negative number")
            
        # Premium type must be either 'above' or 'below'
        if premium_type not in ["above", "below"]:
            raise ValueError("Premium type must be either 'above' or 'below'")
            
        if not isinstance(amount_paid_last_bill, (int, float)) or amount_paid_last_bill < 0:
            raise ValueError("Amount paid last bill must be a non-negative number")
            
        if not isinstance(is_first_bill, bool):
            raise ValueError("is_first_bill must be a boolean value")
            
        # Only validate required fields
        required_user_inputs = ["start_date", "completion_date", "work_order_amount"]
        for field in required_user_inputs:
            if field not in user_inputs or not user_inputs[field]:
                raise ValueError(f"Missing required field: {field}")

        # Set default empty values for optional fields if not provided
        optional_fields = {
            "work_name": "",
            "agreement_no": "",
            "bill_serial": "",
            "work_order_ref": ""
        }
        for field, default_value in optional_fields.items():
            if field not in user_inputs:
                user_inputs[field] = default_value

        # Initialize data structures
        first_page_data = {
            "header": [],
            "items": [],  # Initialize as empty list
            "totals": {}
        }
        last_page_data = {"payable_amount": 0, "amount_words": ""}
        deviation_data = {
            "items": [],  # Initialize as empty list
            "summary": {},
            "header": []
        }
        extra_items_data = {"items": []}
        note_sheet_data = {"notes": []}

        # Header (A1:G19)
        header_data = ws_wo.iloc[:19, :7].replace(pd.NA, "").values.tolist()
        
        # Format dates in header_data
        for i in range(len(header_data)):
            for j in range(len(header_data[i])):
                val = header_data[i][j]
                if isinstance(val, (pd.Timestamp, datetime, date)):
                    header_data[i][j] = val.strftime("%d-%m-%Y")

        # Assign header to both first_page_data and deviation_data
        first_page_data["header"] = header_data
        deviation_data["header"] = header_data

        # Process Work Order items
        last_row_wo = ws_wo.shape[0]
        for i in range(21, last_row_wo):
            qty_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else None
            rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None

            qty = 0
            if isinstance(qty_raw, (int, float)):
                qty = float(qty_raw)
            elif isinstance(qty_raw, str):
                cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
                try:
                    qty = float(cleaned_qty)
                except ValueError:
                    st.warning(f"Skipping invalid quantity at Bill Quantity row {i+1}: '{qty_raw}'")
                    continue

            rate = 0
            if isinstance(rate_raw, (int, float)):
                rate = float(rate_raw)
            elif isinstance(rate_raw, str):
                cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
                try:
                    rate = float(cleaned_rate)
                except ValueError:
                    st.warning(f"Skipping invalid rate at Work Order row {i+1}: '{rate_raw}'")
                    continue

            item = {
                "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
                "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
                "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
                "quantity": qty,
                "rate": rate,
                "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
                "amount": round(qty * rate) if qty and rate else 0,
                "is_divider": False
            }
            first_page_data["items"].append(item)

        # Add Extra Items divider
        first_page_data["items"].append({
            "description": "Extra Items (With Premium)",
            "bold": True,
            "underline": True,
            "amount": 0,
            "quantity": 0,
            "rate": 0,
            "serial_no": "",
            "unit": "",
            "remark": "",
            "is_divider": True
        })

        # Process Extra Items
        last_row_extra = ws_extra.shape[0]
        for j in range(6, last_row_extra):
            qty_raw = ws_extra.iloc[j, 3] if pd.notnull(ws_extra.iloc[j, 3]) else None
            rate_raw = ws_extra.iloc[j, 5] if pd.notnull(ws_extra.iloc[j, 5]) else None

            qty = 0
            if isinstance(qty_raw, (int, float)):
                qty = float(qty_raw)
            elif isinstance(qty_raw, str):
                cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
                try:
                    qty = float(cleaned_qty)
                except ValueError:
                    st.warning(f"Skipping invalid quantity at Extra Items row {j+1}: '{qty_raw}'")
                    continue

            rate = 0
            if isinstance(rate_raw, (int, float)):
                rate = float(rate_raw)
            elif isinstance(rate_raw, str):
                cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
                try:
                    rate = float(cleaned_rate)
                except ValueError:
                    st.warning(f"Skipping invalid rate at Extra Items row {j+1}: '{rate_raw}'")
                    continue

            item = {
                "serial_no": str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else "",
                "description": str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else "",
                "unit": str(ws_extra.iloc[j, 4]) if pd.notnull(ws_extra.iloc[j, 4]) else "",
                "quantity": qty,
                "rate": rate,
                "remark": str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else "",
                "amount": round(qty * rate) if qty and rate else 0,
                "is_divider": False
            }
            first_page_data["items"].append(item)
            extra_items_data["items"].append(item.copy())

        # Calculate totals
        data_items = [item for item in first_page_data["items"] if not item.get("is_divider", False)]
        total_amount = round(sum(item.get("amount", 0) for item in data_items))
        premium_amount = round(total_amount * (premium_percent / 100))
        payable_amount = round(total_amount + premium_amount)

        # Initialize last_page_data with amount_words
        last_page_data = {
            "payable_amount": payable_amount,
            "amount_words": number_to_words(payable_amount)
        }

        first_page_data["totals"] = {
            "grand_total": total_amount,
            "premium": {
                "percent": premium_percent,
                "amount": premium_amount
            },
            "payable": payable_amount
        }

        # Calculate Certificate III data
        certificate_iii_data = {
            "payable_amount": payable_amount,
            "total_123": total_amount,
            "balance_4_minus_5": payable_amount,
            "amount_paid_last_bill": int(amount_paid_last_bill),
            "payment_now": payable_amount,
            "by_cheque": payable_amount,
            "cheque_amount_words": number_to_words(payable_amount),
            "certificate_items": [
                {"name": "Total value of work", "percentage": "100%", "value": total_amount},
                {"name": "Less: Amount Paid Last Bill", "percentage": "-", "value": int(amount_paid_last_bill)},
                {"name": "Net Payable", "percentage": "-", "value": payable_amount}
            ],
            "total_recovery": 0,  # Add logic for recovery items if needed
            "current_date": datetime.now().strftime("%d-%m-%Y")
        }

        # First Page
        first_page_data["header"] = [
            ["Start Date:", user_inputs.get("start_date", "")],
            ["Completion Date:", user_inputs.get("completion_date", "")],
            ["Actual Completion Date:", user_inputs.get("actual_completion_date", "")],
            ["Order Date:", user_inputs.get("order_date", "")],
            ["Contractor Name:", user_inputs.get("contractor_name", "")],
            ["Work Name:", user_inputs.get("work_name", "")],
            ["Bill Serial:", user_inputs.get("bill_serial", "")],
            ["Agreement No:", user_inputs.get("agreement_no", "")],
            ["Work Order Ref:", user_inputs.get("work_order_ref", "")],
            ["Work Order Amount:", user_inputs.get("work_order_amount", "")],
            ["Premium Percent:", premium_percent],
            ["Amount Paid Last Bill:", amount_paid_last_bill],
            ["Bill Type:", user_inputs.get("bill_type", "")],
            ["Bill Number:", user_inputs.get("bill_number", "")],
            ["Last Bill Reference:", user_inputs.get("last_bill_reference", "")]
        ]

        # Last Page
        last_page_data = {
            "payable_amount": payable_amount,
            "amount_words": number_to_words(payable_amount),
            "current_date": datetime.now().strftime("%d-%m-%Y")
        }

        # Deviation Statement
        deviation_data = {
            "items": [],  # Will be populated with bill items
            "summary": {},
            "current_date": datetime.now().strftime("%d-%m-%Y")
        }

        # Process deviation items
        deviation_items = []
        try:
            for i in range(21, ws_wo.shape[0]):
                if i < ws_bq.shape[0]:
                    try:
                        # Get values safely using iloc with bounds checking
                        qty_wo_val = ws_wo.iloc[i, 3] if 3 < ws_wo.shape[1] else None
                        rate_val = ws_wo.iloc[i, 4] if 4 < ws_wo.shape[1] else None
                        qty_bill_val = ws_bq.iloc[i, 3] if 3 < ws_bq.shape[1] else None
                        
                        # Convert and handle null values
                        qty_wo = round(float(qty_wo_val), 2) if pd.notnull(qty_wo_val) else 0
                        rate = float(rate_val) if pd.notnull(rate_val) else 0
                        qty_bill = round(float(qty_bill_val), 2) if pd.notnull(qty_bill_val) else 0

                        # Calculate amounts as integers
                        amt_wo = int(round(qty_wo * rate))
                        amt_bill = int(round(qty_bill * rate))
                        excess_qty = round(qty_bill - qty_wo, 2) if qty_bill > qty_wo else 0
                        excess_amt = int(round(excess_qty * rate)) if excess_qty > 0 else 0
                        saving_qty = round(qty_wo - qty_bill, 2) if qty_bill < qty_wo else 0
                        saving_amt = int(round(saving_qty * rate)) if saving_qty > 0 else 0

                        # Get text values safely with bounds checking
                        serial_no = str(ws_wo.iloc[i, 0]) if 0 < ws_wo.shape[1] and pd.notnull(ws_wo.iloc[i, 0]) else ""
                        description = str(ws_wo.iloc[i, 1]) if 1 < ws_wo.shape[1] and pd.notnull(ws_wo.iloc[i, 1]) else ""
                        unit = str(ws_wo.iloc[i, 2]) if 2 < ws_wo.shape[1] and pd.notnull(ws_wo.iloc[i, 2]) else ""

                        item = {
                            "serial_no": serial_no,
                            "description": description,
                            "unit": unit,
                            "qty_wo": qty_wo,
                            "rate": rate,
                            "amt_wo": amt_wo,
                            "qty_bill": qty_bill,
                            "amt_bill": amt_bill,
                            "excess_qty": excess_qty if excess_qty > 0 else "",
                            "excess_amt": excess_amt if excess_amt > 0 else "",
                            "saving_qty": saving_qty if saving_qty > 0 else "",
                            "saving_amt": saving_amt if saving_amt > 0 else ""
                        }
                        deviation_items.append(item)
                    except (ValueError, TypeError, IndexError):
                        continue  # Skip this row if any conversion fails
        except Exception as e:
            raise ValueError(f"Error processing deviation items: {str(e)}")

        # Calculate deviation summary
        try:
            work_order_total = sum(item["amt_wo"] for item in deviation_items)
            executed_total = sum(item["amt_bill"] for item in deviation_items)
            overall_excess = sum(item["excess_amt"] for item in deviation_items if isinstance(item["excess_amt"], (int, float)))
            overall_saving = sum(item["saving_amt"] for item in deviation_items if isinstance(item["saving_amt"], (int, float)))
        except Exception as e:
            raise ValueError(f"Error calculating deviation summary: {str(e)}")

        # Calculate tender premium with 2 decimal places
        tender_premium_f = round(work_order_total * (premium_percent / 100), 2) if premium_type == "above" else -round(work_order_total * (premium_percent / 100), 2)
        tender_premium_h = round(executed_total * (premium_percent / 100), 2) if premium_type == "above" else -round(executed_total * (premium_percent / 100), 2)
        tender_premium_j = round(overall_excess * (premium_percent / 100), 2) if premium_type == "above" else -round(overall_excess * (premium_percent / 100), 2)
        tender_premium_l = round(overall_saving * (premium_percent / 100), 2) if premium_type == "above" else -round(overall_saving * (premium_percent / 100), 2)

        # Calculate grand totals as integers
        grand_total_f = int(round(work_order_total + tender_premium_f))
        grand_total_h = int(round(executed_total + tender_premium_h))
        grand_total_j = int(round(overall_excess + tender_premium_j))
        grand_total_l = int(round(overall_saving + tender_premium_l))

        net_difference = grand_total_j - grand_total_l
        net_difference_percent = (net_difference / work_order_total * 100) if work_order_total > 0 else 0

        deviation_data["summary"] = {
            "work_order_total": work_order_total,
            "executed_total": executed_total,
            "overall_excess": overall_excess,
            "overall_saving": overall_saving,
            "premium": {
                "percent": premium_percent / 100,
                "type": premium_type
            },
            "tender_premium_f": tender_premium_f,
            "tender_premium_h": tender_premium_h,
            "tender_premium_j": tender_premium_j,
            "tender_premium_l": tender_premium_l,
            "grand_total_f": grand_total_f,
            "grand_total_h": grand_total_h,
            "grand_total_j": grand_total_j,
            "grand_total_l": grand_total_l,
            "net_difference": net_difference,
            "net_difference_percent": net_difference_percent
        }

        # Note Sheet
        note_sheet_data = {
            "notes": generate_bill_notes(payable_amount, user_inputs.get("work_order_amount", 0), sum(item.get("amount", 0) for item in extra_items_data["items"])),
            "current_date": datetime.now().strftime("%d-%m-%Y")
        }

        return first_page_data, last_page_data, deviation_data, extra_items_data["items"], note_sheet_data, certificate_iii_data

    except Exception as e:
        raise Exception(f"Error processing bill data: {str(e)}")

def generate_bill_notes(payable_amount, work_order_amount, extra_item_amount):
    percentage_work_done = (payable_amount / work_order_amount * 100) if work_order_amount > 0 else 0
    serial_number = 1
    note = []

    note.append(f"{serial_number}. The work has been completed {percentage_work_done:.2f}% of the Work Order Amount.")
    serial_number += 1

    if percentage_work_done < 90:
        note.append(f"{serial_number}. The execution of work at final stage is less than 90% of the Work Order Amount, the Requisite Deviation Statement is enclosed to observe check on unuseful expenditure. Approval of the Deviation is having jurisdiction under this office.")
        serial_number += 1
    elif percentage_work_done > 100 and percentage_work_done <= 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed. The Overall Excess is less than or equal to 5% and is having approval jurisdiction under this office.")
        serial_number += 1
    elif percentage_work_done > 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed. The Overall Excess is more than 5% and Approval of the Deviation Case is required from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
        serial_number += 1

    note.append(f"{serial_number}. Quality Control (QC) test reports attached.")
    serial_number += 1

    if extra_item_amount > 0:
        extra_item_percentage = (extra_item_amount / work_order_amount * 100) if work_order_amount > 0 else 0
        if extra_item_percentage > 5:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}. which is {extra_item_percentage:.2f}% of the Work Order Amount; exceed 5%, require approval from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
        else:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}. which is {extra_item_percentage:.2f}% of the Work Order Amount; under 5%, approval of the same is to be granted by this office.")
        serial_number += 1

    note.append(f"{serial_number}. Please peruse above details for necessary decision-making.")
    note.append("")
    note.append("                                Premlata Jain")
    note.append("                               AAO- As Auditor")

    return {"notes": note}

def merge_pdfs(pdf_files, output_file):
    merger = PdfMerger()
    for pdf in pdf_files:
        merger.append(pdf)
    merger.write(output_file)
    merger.close()

def create_word_doc(sheet_name, data, doc_path):
    doc = Document()
    if sheet_name == "First Page":
        table = doc.add_table(rows=len(data["items"]) + 3, cols=9)
        table.style = "Table Grid"
        for i, item in enumerate(data["items"]):
            row = table.rows[i]
            row.cells[0].text = str(item.get("unit", ""))
            row.cells[2].text = str(item.get("quantity", ""))
            row.cells[3].text = str(item.get("serial_no", ""))
            row.cells[4].text = str(item.get("description", ""))
            row.cells[5].text = str(item.get("rate", ""))
            row.cells[6].text = str(item.get("amount", ""))
            row.cells[8].text = str(item.get("remark", ""))
            if item.get("bold"):
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].bold = True
        # Totals
        row = table.rows[-3]
        row.cells[4].text = "Grand Total"
        row.cells[6].text = str(data["totals"]["grand_total"])
        row = table.rows[-2]
    elif sheet_name == "Last Page":
        doc.add_paragraph(f"Payable Amount: {data['payable_amount']}")
        doc.add_paragraph(f"Total in Words: {data['amount_words']}")
    elif sheet_name == "Extra Items":
        table = doc.add_table(rows=len(data["items"]) + 1, cols=7)
        table.style = "Table Grid"
        headers = ["Serial No.", "Remark", "Description", "Quantity", "Unit", "Rate", "Amount"]
        for j, header in enumerate(headers):
            table.rows[0].cells[j].text = header
        for i, item in enumerate(data["items"]):
            row = table.rows[i + 1]
            row.cells[0].text = str(item["serial_no"])
            row.cells[1].text = str(item["remark"])
            row.cells[2].text = str(item["description"])
            row.cells[3].text = str(item["quantity"])
            row.cells[4].text = str(item["unit"])
            row.cells[5].text = str(item["rate"])
            row.cells[6].text = str(item["amount"])
    elif sheet_name == "Deviation Statement":
        table = doc.add_table(rows=len(data["items"]) + 5, cols=12)
        table.style = "Table Grid"
        headers = ["Serial No.", "Description", "Unit", "Qty WO", "Rate", "Amt WO", "Qty Bill", "Amt Bill", "Excess Qty", "Excess Amt", "Saving Qty", "Saving Amt"]
        for j, header in enumerate(headers):
            table.rows[0].cells[j].text = header
        for i, item in enumerate(data["items"]):
            row = table.rows[i + 1]
            row.cells[0].text = str(item["serial_no"])
            row.cells[1].text = str(item["description"])
            row.cells[2].text = str(item["unit"])
            row.cells[3].text = str(item["qty_wo"])
            row.cells[4].text = str(item["rate"])
            row.cells[5].text = str(item["amt_wo"])
            row.cells[6].text = str(item["qty_bill"])
            row.cells[7].text = str(item["amt_bill"])
            row.cells[8].text = str(item["excess_qty"])
            row.cells[9].text = str(item["excess_amt"])
            row.cells[10].text = str(item["saving_qty"])
            row.cells[11].text = str(item["saving_amt"])
        # Summary
        row = table.rows[-4]
        row.cells[1].text = "Grand Total"
        row.cells[5].text = str(data["summary"]["work_order_total"])
        row.cells[7].text = str(data["summary"]["executed_total"])
        row.cells[9].text = str(data["summary"]["overall_excess"])
        row.cells[11].text = str(data["summary"]["overall_saving"])
        row = table.rows[-3]
        row.cells[1].text = f"Add Tender Premium ({data['summary']['premium']['percent']:.2%} {data['summary']['premium']['type']})"
        row.cells[5].text = str(data["summary"]["tender_premium_f"])
        row.cells[7].text = str(data["summary"]["tender_premium_h"])
        row.cells[9].text = str(data["summary"]["tender_premium_j"])
        row.cells[11].text = str(data["summary"]["tender_premium_l"])
        row = table.rows[-2]
        row.cells[1].text = "Grand Total including Tender Premium"
        row.cells[5].text = str(data["summary"]["grand_total_f"])
        row.cells[7].text = str(data["summary"]["grand_total_h"])
        row.cells[9].text = str(data["summary"]["grand_total_j"])
        row.cells[11].text = str(data["summary"]["grand_total_l"])
        row = table.rows[-1]
        if data["summary"]["net_difference"] > 0:
            row.cells[1].text = "Overall Excess With Respect to the Work Order Amount Rs."
            row.cells[7].text = str(round(data["summary"]["net_difference"]))
            row.cells[9].text = f"{data['summary']['net_difference_percent']:2f}%"
        else:
            row.cells[1].text = "Overall Saving With Respect to the Work Order Amount Rs."
            row.cells[7].text = str(round(-data["summary"]["net_difference"]))
            row.cells[9].text = f"{data['summary']['net_difference_percent']:2f}%"
    elif sheet_name == "Note Sheet":
        for note in data["notes"]:
            p = doc.add_paragraph(note)
            p.runs[0].font.name = "Arial Rounded MT Bold"
    doc.save(doc_path)

# Remove @lru_cache decorator from functions that use dictionaries
def get_first_page_html(data):
    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            table {
                width: 100%;
                border-collapse: collapse;
                font-family: Arial, sans-serif;
                font-size: 12px;
            }
            th, td {
                border: 1px solid black;
                padding: 5px;
                text-align: center;
            }
            .bold {
                font-weight: bold;
            }
            .underline {
                text-decoration: underline;
            }
            .header {
                font-size: 14px;
                font-weight: bold;
                margin-bottom: 10px;
            }
            .footer {
                margin-top: 20px;
                text-align: right;
            }
        </style>
    </head>
    <body>
        <div class="header">
            <p>STATE OF RAJASTHAN</p>
            <p>DEPARTMENT OF PUBLIC WORKS</p>
            <p>Udaipur Division</p>
        </div>

        <table>
            <tr class="bold">
                <th>Unit</th>
                <th>Quantity</th>
                <th>Serial No.</th>
                <th>Description</th>
                <th>Rate</th>
                <th>Amount</th>
                <th>Remark</th>
            </tr>
    """
    
    for item in data["items"]:
        html += "<tr>"
        html += f"<td>{item.get('unit', '')}</td>"
        html += f"<td>{item.get('quantity', '')}</td>"
        html += f"<td>{item.get('serial_no', '')}</td>"
        html += f'<td class="left-align">{item.get("description", "")}</td>'
        html += f"<td>{item.get('rate', '')}</td>"
        html += f"<td>{item.get('amount', '')}</td>"
        html += f'<td class="left-align">{item.get("remark", "")}</td>'
        html += "</tr>"
    
    # Totals
    html += f"""
        <tr class="bold">
            <td colspan="5" class="right-align">Grand Total</td>
            <td>{data["totals"]["grand_total"]}</td>
            <td></td>
        </tr>
    """
    
    html += """
        </table>

        <div class="footer">
            <p>Prepared by:</p>
            <p>Checked by:</p>
            <p>Verified by:</p>
        </div>
    </body>
    </html>
    """
    return html

def get_last_page_html(data):
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                font-size: 14px;
                padding: 20px;
            }}
            .bold {{
                font-weight: bold;
            }}
        </style>
    </head>
    <body>
        <p class="bold">Payable Amount: {data["payable_amount"]}</p>
        <p class="bold">Total in Words: {data["amount_words"]}</p>
    </body>
    </html>
    """
    return html

def get_extra_items_html(data):
    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            table {
                width: 100%;
                border-collapse: collapse;
                font-family: Arial, sans-serif;
                font-size: 12px;
            }
            th, td {
                border: 1px solid black;
                padding: 5px;
                text-align: center;
            }
            .bold {
                font-weight: bold;
            }
            .underline {
                text-decoration: underline;
            }
            .header {
                font-size: 14px;
                font-weight: bold;
                margin-bottom: 10px;
            }
            .footer {
                margin-top: 20px;
                text-align: right;
            }
        </style>
    </head>
    <body>
        <div class="header">
            <p>STATE OF RAJASTHAN</p>
            <p>DEPARTMENT OF PUBLIC WORKS</p>
            <p>Udaipur Division</p>
        </div>

        <table>
            <tr class="bold">
                <th>Serial No.</th>
                <th>Remark</th>
                <th>Description</th>
                <th>Quantity</th>
                <th>Unit</th>
                <th>Rate</th>
                <th>Amount</th>
            </tr>
    """
    
    for item in data["items"]:
        html += "<tr>"
        html += f"<td>{item['serial_no']}</td>"
        html += f'<td class="left-align">{item["remark"]}</td>'
        html += f'<td class="left-align">{item["description"]}</td>'
        html += f"<td>{item['quantity']}</td>"
        html += f"<td>{item['unit']}</td>"
        html += f"<td>{item['rate']}</td>"
        html += f"<td>{item['amount']}</td>"
        html += "</tr>"
    
    html += """
        </table>

        <div class="footer">
            <p>Prepared by:</p>
            <p>Checked by:</p>
            <p>Verified by:</p>
        </div>
    </body>
    </html>
    """
    return html

def get_deviation_statement_html(data):
    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            table {
                width: 100%;
                border-collapse: collapse;
                font-family: Arial, sans-serif;
                font-size: 12px;
            }
            th, td {
                border: 1px solid black;
                padding: 5px;
                text-align: center;
            }
            .bold {
                font-weight: bold;
            }
            .underline {
                text-decoration: underline;
            }
            .header {
                font-size: 14px;
                font-weight: bold;
                margin-bottom: 10px;
            }
            .footer {
                margin-top: 20px;
                text-align: right;
            }
        </style>
    </head>
    <body>
        <div class="header">
            <p>STATE OF RAJASTHAN</p>
            <p>DEPARTMENT OF PUBLIC WORKS</p>
            <p>Udaipur Division</p>
        </div>

        <table>
            <tr class="bold">
                <th>Serial No.</th>
                <th>Description</th>
                <th>Unit</th>
                <th>Qty WO</th>
                <th>Rate</th>
                <th>Amt WO</th>
                <th>Qty Bill</th>
                <th>Amt Bill</th>
                <th>Excess Qty</th>
                <th>Excess Amt</th>
                <th>Saving Qty</th>
                <th>Saving Amt</th>
            </tr>
    """
    
    for item in data["items"]:
        html += "<tr>"
        html += f"<td>{item['serial_no']}</td>"
        html += f"<td>{item['description']}</td>"
        html += f"<td>{item['unit']}</td>"
        html += f"<td>{item['qty_wo']}</td>"
        html += f"<td>{item['rate']}</td>"
        html += f"<td>{item['amt_wo']}</td>"
        html += f"<td>{item['qty_bill']}</td>"
        html += f"<td>{item['amt_bill']}</td>"
        html += f"<td>{item['excess_qty']}</td>"
        html += f"<td>{item['excess_amt']}</td>"
        html += f"<td>{item['saving_qty']}</td>"
        html += f"<td>{item['saving_amt']}</td>"
        html += "</tr>"
    
    # Summary
    html += f"""
        <tr class="bold">
            <td colspan="5" class="right-align">Grand Total</td>
            <td>{data["summary"]["work_order_total"]}</td>
            <td>{data["summary"]["executed_total"]}</td>
            <td></td>
            <td>{data["summary"]["overall_excess"]}</td>
            <td></td>
            <td>{data["summary"]["overall_saving"]}</td>
        </tr>
        <tr class="bold">
            <td colspan="5" class="right-align">Add Tender Premium ({data["summary"]["premium"]["percent"]:.2%} {data["summary"]["premium"]["type"]})</td>
            <td>{data["summary"]["tender_premium_f"]}</td>
            <td>{data["summary"]["tender_premium_h"]}</td>
            <td></td>
            <td>{data["summary"]["tender_premium_j"]}</td>
            <td></td>
            <td>{data["summary"]["tender_premium_l"]}</td>
        </tr>
        <tr class="bold">
            <td colspan="5" class="right-align">Grand Total including Tender Premium</td>
            <td>{data["summary"]["grand_total_f"]}</td>
            <td>{data["summary"]["grand_total_h"]}</td>
            <td></td>
            <td>{data["summary"]["grand_total_j"]}</td>
            <td></td>
            <td>{data["summary"]["grand_total_l"]}</td>
        </tr>
    """
    
    if data["summary"]["net_difference"] > 0:
        html += f"""
        <tr class="bold">
            <td colspan="7" class="right-align">Overall Excess With Respect to the Work Order Amount Rs.</td>
            <td colspan="2">{round(data["summary"]["net_difference"])}</td>
            <td colspan="2">{data["summary"]["net_difference_percent"]:.2f}%</td>
        </tr>
        """
    else:
        html += f"""
        <tr class="bold">
            <td colspan="7" class="right-align">Overall Saving With Respect to the Work Order Amount Rs.</td>
            <td colspan="2">{round(-data["summary"]["net_difference"])}</td>
            <td colspan="2">{data["summary"]["net_difference_percent"]:.2f}%</td>
        </tr>
        """
    
    html += """
        </table>
    </body>
    </html>
    """
    return html

def get_note_sheet_html(data):
    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {
                font-family: Arial, sans-serif;
                font-size: 14px;
                padding: 20px;
            }
            p {
                margin-bottom: 10px;
            }
            .signature {
                margin-top: 50px;
                text-align: right;
            }
        </style>
    </head>
    <body>
    """
    
    for note in data["notes"]:
        html += f"<p>{note}</p>"
    
    html += """
        <div class="signature">
            <p>Signature: ______________________</p>
            <p>Date: {data["current_date"]}</p>
        </div>
    </body>
    </html>
    """
    return html

def get_certificate_iii_html(data):
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                font-size: 14px;
                padding: 20px;
            }}
            .bold {{
                font-weight: bold;
            }}
        </style>
    </head>
    <body>
        <p class="bold">Payable Amount: {data["payable_amount"]}</p>
        <p class="bold">Total in Words: {data["amount_words"]}</p>
    </body>
    </html>
    """
    return html

def generate_pdf_files(html_files, output_dir):
    """Generate PDF files in parallel"""
    try:
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = []
            for html_file in html_files:
                pdf_file = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(html_file))[0]}.pdf")
                futures.append(executor.submit(pdfkit.from_file, html_file, pdf_file, configuration=config))
            
            # Wait for all PDFs to be generated
            concurrent.futures.wait(futures)
            
            # Check for any errors
            for future in futures:
                try:
                    future.result()
                except Exception as e:
                    logger.error(f"Error generating PDF: {str(e)}")
                    raise Exception(f"Error generating PDF: {str(e)}")
                    
            logger.info("All PDF files generated successfully")
    except Exception as e:
        logger.error(f"Error in PDF generation process: {str(e)}")
        raise

def sanitize_input(value: Any, field_type: str) -> Any:
    """
    Sanitize user input based on field type.
    
    Args:
        value: The input value to sanitize
        field_type: The type of field ('date', 'number', 'text', 'boolean')
        
    Returns:
        The sanitized value
        
    Raises:
        ValueError: If the value cannot be sanitized
    """
    try:
        if field_type == 'date':
            if isinstance(value, str):
                return datetime.strptime(value, "%Y-%m-%d").date()
            return value
        elif field_type == 'number':
            if isinstance(value, str):
                return float(value)
            return value
        elif field_type == 'text':
            if isinstance(value, str):
                return value.strip()
            return str(value)
        elif field_type == 'boolean':
            if isinstance(value, str):
                return value.lower() in ('true', '1', 'yes')
            return bool(value)
        else:
            return value
    except Exception as e:
        raise ValueError(f"Error sanitizing {field_type} input: {str(e)}")

def validate_user_inputs(user_inputs: Dict[str, Any]) -> Dict[str, Any]:
    """
    Validate and sanitize user inputs.
    
    Args:
        user_inputs: Dictionary of user inputs
        
    Returns:
        Dictionary of validated and sanitized inputs
        
    Raises:
        ValueError: If validation fails
    """
    try:
        validated_inputs = {}
        
        # Required fields
        required_fields = {
            'start_date': 'date',
            'completion_date': 'date',
            'work_order_amount': 'number',
            'premium_percent': 'number',
            'premium_type': 'text',
            'amount_paid_last_bill': 'number',
            'is_first_bill': 'boolean'
        }
        
        for field, field_type in required_fields.items():
            if field not in user_inputs:
                raise ValueError(f"Missing required field: {field}")
            validated_inputs[field] = sanitize_input(user_inputs[field], field_type)
            
        # Optional fields
        optional_fields = {
            'work_name': 'text',
            'agreement_no': 'text',
            'bill_serial': 'text',
            'work_order_ref': 'text'
        }
        
        for field, field_type in optional_fields.items():
            if field in user_inputs:
                validated_inputs[field] = sanitize_input(user_inputs[field], field_type)
            else:
                validated_inputs[field] = ''
                
        # Additional validation
        if validated_inputs['start_date'] > validated_inputs['completion_date']:
            raise ValueError("Start date cannot be after completion date")
            
        if validated_inputs['work_order_amount'] <= 0:
            raise ValueError("Work order amount must be positive")
            
        if validated_inputs['premium_percent'] < 0 or validated_inputs['premium_percent'] > 100:
            raise ValueError("Premium percentage must be between 0 and 100")
            
        if validated_inputs['premium_type'] not in ['above', 'below']:
            raise ValueError("Premium type must be either 'above' or 'below'")
            
        if validated_inputs['amount_paid_last_bill'] < 0:
            raise ValueError("Amount paid last bill cannot be negative")
            
        return validated_inputs
        
    except Exception as e:
        raise ValueError(f"Error validating user inputs: {str(e)}")

def main():
    st.title("Contractor Bill Generator")
    
    try:
        # Add custom CSS for styling
        st.markdown("""
        <style>
            .required-field::after {
                content: " *";
                color: red;
            }
            .stButton > button {
                background-color: #4CAF50;
                color: white;
                padding: 10px 20px;
                border: none;
                border-radius: 4px;
                cursor: pointer;
                font-size: 16px;
            }
            .stButton > button:hover {
                background-color: #45a049;
            }
            .stFileUploader > div > div {
                background-color: #f5f5f5;
                padding: 10px;
                border-radius: 4px;
                border: 1px solid #ddd;
            }
            .stFileUploader > div > div:hover {
                background-color: #e0e0e0;
            }
            .stFormSubmitButton > button {
                background-color: #2196F3 !important;
                color: white !important;
            }
            .stFormSubmitButton > button:hover {
                background-color: #1976D2 !important;
            }
        </style>
        """, unsafe_allow_html=True)

        # Create a form for user inputs
        with st.form("bill_input_form"):
            st.subheader("Bill Details")
            
            # Basic Information
            col1, col2 = st.columns(2)
            with col1:
                st.markdown('<p class="required-field">Start Date</p>', unsafe_allow_html=True)
                start_date = st.date_input("", key="start_date")
                work_name = st.text_input("Work Name")
                agreement_no = st.text_input("Agreement Number")
            with col2:
                st.markdown('<p class="required-field">Completion Date</p>', unsafe_allow_html=True)
                completion_date = st.date_input("", key="completion_date")
                bill_serial = st.text_input("Bill Serial")
                work_order_ref = st.text_input("Work Order Reference")
                st.markdown('<p class="required-field">Work Order Amount</p>', unsafe_allow_html=True)
                work_order_amount = st.number_input(
                    "",
                    min_value=0,
                    value=0,
                    step=1,
                    help="Enter the total work order amount (rounded to nearest integer, required)"
                )

            # Premium Information
            st.subheader("Premium Information")
            st.markdown('<p class="required-field">Premium Type</p>', unsafe_allow_html=True)
            premium_type = st.radio(
                "",
                ["Above", "Below"],
                horizontal=True,
                help="Select whether the premium is above or below the base amount"
            ).lower()
            
            st.markdown('<p class="required-field">Premium Percentage</p>', unsafe_allow_html=True)
            premium_percent = st.number_input(
                "",
                min_value=0.0,
                max_value=100.0,
                value=0.0,
                step=0.01,
                format="%.2f",
                help=f"Enter the premium percentage to be added {'above' if premium_type == 'above' else 'below'} the base amount (2 decimal places)"
            )
            
            # Payment Information
            st.subheader("Payment Information")
            st.markdown('<p class="required-field">Amount Paid in Last Bill</p>', unsafe_allow_html=True)
            amount_paid_last_bill = st.number_input(
                "",
                min_value=0,
                value=0
            )
            st.markdown('<p class="required-field">Is this the first bill?</p>', unsafe_allow_html=True)
            is_first_bill = st.checkbox("")

            # File upload
            st.subheader("Upload Excel File")
            st.markdown('<p class="required-field">Excel File</p>', unsafe_allow_html=True)
            uploaded_file = st.file_uploader(
                "",
                type=["xlsx", "xls"],
                help="Upload an Excel file containing Work Order, Bill Quantity, and Extra Items sheets"
            )

            # Submit button
            submit_button = st.form_submit_button("Generate Bill")

        # Process the bill when submitted
        if submit_button:
            if uploaded_file is None:
                st.error("Please upload an Excel file first")
                return

            try:
                # Validate and sanitize user inputs
                user_inputs = validate_user_inputs({
                    "start_date": start_date,
                    "completion_date": completion_date,
                    "work_name": work_name or "",  # Make optional
                    "bill_serial": bill_serial or "",  # Make optional
                    "agreement_no": agreement_no or "",  # Make optional
                    "work_order_ref": work_order_ref or "",  # Make optional
                    "work_order_amount": work_order_amount,
                    "premium_percent": premium_percent,
                    "premium_type": premium_type,
                    "amount_paid_last_bill": amount_paid_last_bill,
                    "is_first_bill": is_first_bill
                })
                
                # Create a new temporary directory for this run
                temp_dir = tempfile.mkdtemp()
                logger.info(f"Created temporary directory: {temp_dir}")
                
                # Read and validate the uploaded file
                with pd.ExcelFile(uploaded_file) as xls:
                    validate_excel_sheets(xls)
                    ws_wo = pd.read_excel(xls, "Work Order", header=None)
                    ws_bq = pd.read_excel(xls, "Bill Quantity", header=None)
                    ws_extra = pd.read_excel(xls, "Extra Items", header=None)

                # Process the bill
                first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data, certificate_iii_data = process_bill(
                    ws_wo,
                    ws_bq,
                    ws_extra,
                    user_inputs["premium_percent"],
                    user_inputs["premium_type"],
                    user_inputs["amount_paid_last_bill"],
                    user_inputs["is_first_bill"],
                    user_inputs
                )

                # Generate PDFs
                pdf_files = []
                word_files = []
                
                # Mapping of sheet names to their HTML generation functions
                html_generators = {
                    "First Page": get_first_page_html,
                    "Last Page": get_last_page_html,
                    "Extra Items": get_extra_items_html,
                    "Deviation Statement": get_deviation_statement_html,
                    "Note Sheet": get_note_sheet_html,
                    "Certificate III": get_certificate_iii_html
                }

                # Generate HTML files
                html_files = []
                for sheet_name, data in [
                    ("First Page", first_page_data),
                    ("Last Page", last_page_data),
                    ("Extra Items", {"items": extra_items_data}),
                    ("Deviation Statement", deviation_data),
                    ("Note Sheet", note_sheet_data),
                    ("Certificate III", certificate_iii_data)
                ]:
                    html_file = os.path.join(temp_dir, f"{sheet_name.replace(' ', '_')}.html")
                    html_generator = html_generators.get(sheet_name)
                    if html_generator:
                        html_content = html_generator(data)
                        with open(html_file, "w", encoding="utf-8") as f:
                            f.write(html_content)
                        html_files.append(html_file)
                    else:
                        logger.error(f"No HTML generator found for sheet: {sheet_name}")
                        raise ValueError(f"No HTML generator found for sheet: {sheet_name}")

                # Generate PDFs in parallel
                generate_pdf_files(html_files, temp_dir)
                
                # Merge PDFs
                merger = PdfMerger()
                for pdf_file in [os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(f))[0]}.pdf") for f in html_files]:
                    merger.append(pdf_file)
                pdf_output = os.path.join(temp_dir, "output.pdf")
                merger.write(pdf_output)
                merger.close()

                # Generate Word documents
                for sheet_name, data in [
                    ("First Page", first_page_data),
                    ("Last Page", last_page_data),
                    ("Extra Items", {"items": extra_items_data}),
                    ("Deviation Statement", deviation_data),
                    ("Note Sheet", note_sheet_data)
                ]:
                    doc_path = os.path.join(temp_dir, f"{sheet_name.replace(' ', '_')}.docx")
                    create_word_doc(sheet_name, data, doc_path)
                    word_files.append(doc_path)

                # Create ZIP file
                zip_path = os.path.join(temp_dir, "output.zip")
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                    zipf.write(pdf_output, os.path.basename(pdf_output))
                    for word_file in word_files:
                        zipf.write(word_file, os.path.basename(word_file))

                # Provide download link
                with open(zip_path, "rb") as f:
                    bytes_data = f.read()
                st.download_button(
                    label="Download Output Files",
                    data=bytes_data,
                    file_name="bill_output.zip",
                    mime="application/zip"
                )
                
                logger.info("Bill generation completed successfully")
                
            except Exception as e:
                logger.error(f"Error processing file: {str(e)}")
                st.error(f"Error processing file: {str(e)}")
                st.stop()
            finally:
                # Clean up temporary files
                if os.path.exists(temp_dir):
                    try:
                        shutil.rmtree(temp_dir)
                        logger.info(f"Cleaned up temporary directory: {temp_dir}")
                    except Exception as e:
                        logger.error(f"Error cleaning up temporary directory: {str(e)}")

    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        st.error(f"Unexpected error: {str(e)}")
        st.stop()

if __name__ == "__main__":
    main()